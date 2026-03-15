#!/usr/bin/env python3
"""
Reference Alignment Tool - 参考文档对齐工具

功能：
1. 读取参考PDF/Word文档
2. 分析文档结构、内容密度、表格数量
3. 生成对齐检查清单
4. 提供内容密度对比报告

使用场景：
- 用户提供了参考报告
- 需要匹配参考报告的深度和风格
"""

import fitz  # PyMuPDF
import re
from pathlib import Path
from typing import Dict, List, Tuple
import json


class ReferenceAnalyzer:
    """参考文档分析器"""

    def __init__(self, reference_path: str):
        """
        初始化分析器

        Args:
            reference_path: 参考文档路径（PDF/Word）
        """
        self.reference_path = Path(reference_path)
        self.analysis_result = {}

        if not self.reference_path.exists():
            raise FileNotFoundError(f"参考文档不存在: {reference_path}")

    def analyze(self) -> Dict:
        """
        分析参考文档

        Returns:
            分析结果字典，包含：
            - page_count: 页数
            - structure: 文档结构（章节列表）
            - tables: 表格数量和类型
            - data_points: 数据点估算
            - data_density: 每页数据点密度
            - confidence_levels: 置信度标注情况
        """
        if self.reference_path.suffix.lower() == '.pdf':
            self._analyze_pdf()
        elif self.reference_path.suffix.lower() in ['.docx', '.doc']:
            self._analyze_word()
        else:
            raise ValueError(f"不支持的文件格式: {self.reference_path.suffix}")

        return self.analysis_result

    def _analyze_pdf(self):
        """分析PDF文档"""
        doc = fitz.open(str(self.reference_path))

        # 基本信息
        self.analysis_result['page_count'] = len(doc)

        # 提取文本内容
        full_text = ""
        for page in doc:
            full_text += page.get_text()

        # 分析结构
        self.analysis_result['structure'] = self._extract_structure(full_text)

        # 统计表格
        self.analysis_result['tables'] = self._count_tables(doc)

        # 估算数据点
        self.analysis_result['data_points'] = self._estimate_data_points(full_text)

        # 计算密度
        if self.analysis_result['page_count'] > 0:
            self.analysis_result['data_density'] = (
                self.analysis_result['data_points'] / self.analysis_result['page_count']
            )
        else:
            self.analysis_result['data_density'] = 0

        # 检测置信度标注
        self.analysis_result['confidence_levels'] = self._check_confidence_levels(full_text)

        doc.close()

    def _analyze_word(self):
        """分析Word文档"""
        # 简化版本，使用docx库
        try:
            from docx import Document
            doc = Document(str(self.reference_path))

            # 基本信息
            self.analysis_result['page_count'] = len(doc.sections)  # 估算

            # 提取文本
            full_text = "\n".join([p.text for p in doc.paragraphs])

            # 分析结构
            self.analysis_result['structure'] = self._extract_structure(full_text)

            # 统计表格
            self.analysis_result['tables'] = len(doc.tables)

            # 估算数据点
            self.analysis_result['data_points'] = self._estimate_data_points(full_text)

            # 计算密度
            if self.analysis_result['page_count'] > 0:
                self.analysis_result['data_density'] = (
                    self.analysis_result['data_points'] / self.analysis_result['page_count']
                )

            # 检测置信度标注
            self.analysis_result['confidence_levels'] = self._check_confidence_levels(full_text)

        except ImportError:
            raise ImportError("需要安装python-docx库: pip install python-docx")

    def _extract_structure(self, text: str) -> List[Dict]:
        """
        提取文档结构

        Returns:
            章节列表，如：
            [
                {'level': 1, 'title': '第一章', 'page': 1},
                {'level': 2, 'title': '1.1 市场规模', 'page': 2},
            ]
        """
        structure = []

        # 常见章节标题模式
        patterns = [
            r'(第[一二三四五六七八九十]章|Chapter \d+)[\s：:](.+)',
            r'(\d+\.\d+)\s+(.+)',  # 1.1 格式
            r'(##+)\s+(.+)',  # Markdown格式
        ]

        lines = text.split('\n')
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            for pattern in patterns:
                match = re.match(pattern, line)
                if match:
                    title = match.group(2) if len(match.groups()) > 1 else match.group(0)
                    level = len(match.group(1)) if match.group(1).startswith('#') else 1

                    structure.append({
                        'title': title[:100],  # 限制长度
                        'level': level,
                        'line': line_num
                    })
                    break

        return structure

    def _count_tables(self, pdf_doc) -> Dict:
        """
        统计PDF中的表格

        Returns:
            {'total': 总数, 'by_page': {页码: 数量}}
        """
        tables = {'total': 0, 'by_page': {}}

        for page_num, page in enumerate(pdf_doc):
            try:
                # 查找表格（简化方法：检测表格状的文本排列）
                tables_found = page.find_tables()
                if tables_found:
                    count = len(tables_found)
                    tables['total'] += count
                    tables['by_page'][page_num + 1] = count
            except:
                pass

        return tables

    def _estimate_data_points(self, text: str) -> int:
        """
        估算数据点数量

        查找模式：
        - 数字+单位（亿元、%、万人等）
        - 年份数据（2020-2024等）
        - 百分比（XX%）
        """
        patterns = [
            r'\d+\.\d+[亿元万千百十]',  # 金额
            r'\d+\.?\d*%',  # 百分比
            r'20[12]\d',  # 年份
            r'\d{4}',  # 四位数（可能是年份）
        ]

        count = 0
        for pattern in patterns:
            matches = re.findall(pattern, text)
            count += len(matches)

        # 去重（避免重复计数）
        return min(count, len(text) // 100)  # 估算上限

    def _check_confidence_levels(self, text: str) -> Dict:
        """
        检测置信度标注

        Returns:
            {'has_confidence': bool, 'count': 数量}
        """
        patterns = [
            r'置信度[:：]\s*\d+%',
            r'置信度[:：]\s*[高低中]',
            r'\[置信度',
        ]

        count = 0
        for pattern in patterns:
            matches = re.findall(pattern, text)
            count += len(matches)

        return {
            'has_confidence': count > 0,
            'count': count
        }

    def generate_alignment_report(self, target_spec: Dict = None) -> str:
        """
        生成对齐报告

        Args:
            target_spec: 目标规格，如：
                {
                    'target_pages': 50,
                    'target_tables': 35,
                    'target_data_points': 200,
                }

        Returns:
            报告文本
        """
        if not self.analysis_result:
            self.analyze()

        report = []
        report.append("=" * 60)
        report.append("📊 参考文档对齐分析报告")
        report.append("=" * 60)
        report.append("")

        # 基本信息
        report.append("📋 参考文档基本信息")
        report.append(f"  文件: {self.reference_path.name}")
        report.append(f"  页数: {self.analysis_result.get('page_count', 'N/A')}")
        report.append(f"  表格数: {self.analysis_result.get('tables', {}).get('total', 'N/A') if isinstance(self.analysis_result.get('tables'), dict) else self.analysis_result.get('tables', 'N/A')}")
        report.append(f"  数据点估算: {self.analysis_result.get('data_points', 'N/A')}")
        report.append("")

        # 内容密度
        report.append("📈 内容密度分析")
        density = self.analysis_result.get('data_density', 0)
        report.append(f"  每页数据点密度: {density:.1f} 个/页")

        if density > 10:
            level = "高（非常详细）"
        elif density > 6:
            level = "中（适度详细）"
        else:
            level = "低（简要概述）"
        report.append(f"  密度等级: {level}")
        report.append("")

        # 结构分析
        structure = self.analysis_result.get('structure', [])
        if structure:
            report.append("📑 文档结构")
            for item in structure[:15]:  # 显示前15个
                indent = "  " * (item.get('level', 1) - 1)
                report.append(f"{indent}- {item['title']}")
            if len(structure) > 15:
                report.append(f"  ... (共{len(structure)}个章节)")
            report.append("")

        # 置信度标注
        conf = self.analysis_result.get('confidence_levels', {})
        report.append("✅ 数据质量标注")
        if conf.get('has_confidence'):
            report.append(f"  ✓ 有置信度标注 ({conf.get('count', 0)}处)")
        else:
            report.append(f"  ✗ 未发现置信度标注")
        report.append("")

        # 对比目标（如果提供）
        if target_spec:
            report.append("🎯 目标对齐建议")
            report.append("")

            target_pages = target_spec.get('target_pages', 50)
            target_tables = target_spec.get('target_tables', 35)
            target_data_points = target_spec.get('target_data_points', 200)

            # 页数对比
            current_pages = self.analysis_result.get('page_count', 0)
            if current_pages < target_pages:
                report.append(f"  页数: 参考文档{current_pages}页 < 目标{target_pages}页")
                report.append(f"    → 建议增加深度分析章节")
            elif current_pages > target_pages:
                report.append(f"  页数: 参考文档{current_pages}页 > 目标{target_pages}页")
                report.append(f"    → 可适当精简次要内容")
            else:
                report.append(f"  页数: ✓ 匹配 ({current_pages}页)")

            # 表格对比
            current_tables = self.analysis_result.get('tables', {}).get('total', 0) if isinstance(self.analysis_result.get('tables'), dict) else self.analysis_result.get('tables', 0)
            if current_tables < target_tables:
                report.append(f"  表格: 参考文档{current_tables}个 < 目标{target_tables}个")
                report.append(f"    → 建议增加数据表格，特别是:")
                report.append(f"      - 历史数据表（6年+）")
                report.append(f"      - 竞争对比表（8家+）")
                report.append(f"      - 分析框架表（PESTLE/Porter等）")
            else:
                report.append(f"  表格: ✓ 充足 ({current_tables}个)")

            # 数据点对比
            current_density = self.analysis_result.get('data_density', 0)
            target_density = target_data_points / target_pages

            if current_density < target_density:
                report.append(f"  数据密度: 参考文档{current_density:.1f} < 目标{target_density:.1f} 个/页")
                report.append(f"    → 建议每页增加{target_density - current_density:.1f}个数据点")
            else:
                report.append(f"  数据密度: ✓ 达标 ({current_density:.1f}个/页)")

            report.append("")

        # 建议
        report.append("💡 生成建议")
        report.append("")

        based_on_density = self.analysis_result.get('data_density', 0)

        if based_on_density > 10:
            report.append("  参考文档内容密度非常高，建议:")
            report.append("  1. 每章至少包含5-7个数据表格")
            report.append("  2. 每页至少8-12个数据点")
            report.append("  3. 所有数据标注来源和置信度")
            report.append("  4. 深度分析不只数据罗列")
            report.append("  5. 关键章节扩展到8-10页")
        elif based_on_density > 6:
            report.append("  参考文档内容密度中等，建议:")
            report.append("  1. 每章至少包含4-5个数据表格")
            report.append("  2. 每页至少6-8个数据点")
            report.append("  3. 重要数据标注来源")
            report.append("  4. 重点章节详细分析")
        else:
            report.append("  参考文档内容密度较低，建议:")
            report.append("  1. 补充更多数据支撑")
            report.append("  2. 增加表格数量")
            report.append("  3. 扩展分析深度")

        report.append("")
        report.append("=" * 60)

        return "\n".join(report)


def quick_analyze(reference_path: str) -> str:
    """
    快速分析参考文档并生成报告

    Args:
        reference_path: 参考文档路径

    Returns:
        分析报告文本
    """
    try:
        analyzer = ReferenceAnalyzer(reference_path)
        analyzer.analyze()
        return analyzer.generate_alignment_report({
            'target_pages': 50,
            'target_tables': 35,
            'target_data_points': 200,
        })
    except Exception as e:
        return f"❌ 分析失败: {str(e)}"


if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("用法: python reference_alignment.py <参考文档路径>")
        print("示例: python reference_alignment.py ~/Downloads/参考报告.pdf")
        sys.exit(1)

    reference_path = sys.argv[1]
    report = quick_analyze(reference_path)
    print(report)

    # 保存报告
    report_path = Path(reference_path).parent / f"{Path(reference_path).stem}_对齐分析报告.txt"
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(report)

    print(f"\n✅ 分析报告已保存: {report_path}")
