#!/usr/bin/env python3
"""
Complete Markdown to Word Converter for Market Research Reports

Features:
- Full markdown parsing (headers, tables, lists, emphasis)
- Chinese font handling (黑体 for titles, 宋体 for body)
- Confidence level notes with special formatting
- Source citations preservation
- Emoji support in headers

Author: Pseudoctor Research Team
Version: 1.0.0
Date: 2025-01-30
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path


# Color constants
COLOR_BLUE = RGBColor(0, 51, 102)      # Dark blue for titles
COLOR_DARK_GRAY = RGBColor(51, 51, 51) # Dark gray for body
COLOR_GREEN = RGBColor(0, 100, 0)      # Green for positive items
COLOR_RED = RGBColor(200, 0, 0)        # Red for warnings
COLOR_ORANGE = RGBColor(255, 140, 0)   # Orange for caution
COLOR_PURPLE = RGBColor(128, 0, 128)   # Purple for key points


def set_chinese_font(run, font_name='宋体', size=11, bold=False, color=None):
    """
    Set Chinese font for a text run with proper handling

    Args:
        run: The text run to format
        font_name: Font name (黑体 for titles, 宋体 for body)
        size: Font size in points
        bold: Whether text is bold
        color: RGBColor object for text color
    """
    # Set English font
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color

    # Set Chinese font - critical for Chinese characters
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def parse_table(markdown_table):
    """
    Parse markdown table into headers and rows

    Args:
        markdown_table: List of lines representing a markdown table

    Returns:
        tuple: (headers, rows) where headers is list of strings and rows is list of lists
    """
    lines = [line.strip() for line in markdown_table if line.strip()]

    if len(lines) < 2:
        return None, None

    # Parse header
    header_line = lines[0]
    headers = [cell.strip() for cell in re.split(r'\|', header_line)[1:-1]]

    # Skip separator line
    data_lines = lines[2:] if len(lines) > 2 else []

    # Parse data rows
    rows = []
    for line in data_lines:
        cells = [cell.strip() for cell in re.split(r'\|', line)[1:-1]]
        if cells:
            rows.append(cells)

    return headers, rows


def add_formatted_table(doc, headers, rows, style='Light Grid Accent 1'):
    """
    Add a formatted table to the document

    Args:
        doc: Word document object
        headers: List of header strings
        rows: List of row lists
        style: Table style name
    """
    if not headers or not rows:
        return

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = style

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # Format header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '黑体', 10, True, COLOR_BLUE)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add data rows
    for i, row in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row):
            if j < len(row_cells):
                cell = row_cells[j]
                cell.text = cell_data
                # Format cell text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run, '宋体', 10, False, COLOR_DARK_GRAY)


def add_confidence_note(doc, text, confidence, source):
    """
    Add a formatted confidence note with source citation

    Args:
        doc: Word document object
        text: The main text
        confidence: Confidence level (e.g., '85%')
        source: Source description
    """
    p = doc.add_paragraph()
    run1 = p.add_run(f'{text} ')
    set_chinese_font(run1, '宋体', 10, False, COLOR_DARK_GRAY)

    run2 = p.add_run(f'[置信度: {confidence}]')
    set_chinese_font(run2, '宋体', 9, False, COLOR_ORANGE)

    if source:
        run3 = p.add_run(f'\n来源: {source}')
        set_chinese_font(run3, '宋体', 8, False, RGBColor(100, 100, 100))


def add_key_insight_box(doc, title, insights):
    """
    Add a highlighted key insights box

    Args:
        doc: Word document object
        title: Box title
        insights: List of insight strings
    """
    p = doc.add_paragraph()
    title_run = p.add_run(title)
    set_chinese_font(title_run, '黑体', 11, True, COLOR_PURPLE)

    for insight in insights:
        p = doc.add_paragraph(insight, style='List Bullet')
        for run in p.runs:
            set_chinese_font(run, '宋体', 10, False, COLOR_DARK_GRAY)


def add_warning_box(doc, title, warnings):
    """
    Add a warning/risk box

    Args:
        doc: Word document object
        title: Box title
        warnings: List of warning strings
    """
    p = doc.add_paragraph()
    title_run = p.add_run(title)
    set_chinese_font(title_run, '黑体', 11, True, COLOR_RED)

    for warning in warnings:
        p = doc.add_paragraph(warning, style='List Bullet')
        for run in p.runs:
            set_chinese_font(run, '宋体', 10, False, COLOR_DARK_GRAY)


def convert_markdown_to_word(markdown_path, output_path):
    """
    Convert a complete markdown report to Word format

    Args:
        markdown_path: Path to input markdown file
        output_path: Path to output Word document
    """
    # Read markdown file
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    in_code_block = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue

        # Handle tables (multi-line)
        if '|' in line and line.strip().startswith('|'):
            in_table = True
            table_lines.append(line)
            i += 1
            # Continue collecting table lines
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # Parse and add the table
            headers, rows = parse_table(table_lines)
            if headers and rows:
                add_formatted_table(doc, headers, rows)
            table_lines = []
            in_table = False
            continue

        # Handle headers
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            title_text = re.sub(r'^#+\s*', '', line).strip()

            # Remove emojis from title for Word (optional, keeps them for now)
            p = doc.add_paragraph()
            run = p.add_run(title_text)

            # Font size based on level
            sizes = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11}
            size = sizes.get(level, 11)
            set_chinese_font(run, '黑体', size, True, COLOR_BLUE)

            # Set heading style
            if level == 1:
                p.style = 'Heading 1'
            elif level == 2:
                p.style = 'Heading 2'
            elif level == 3:
                p.style = 'Heading 3'
            elif level == 4:
                p.style = 'Heading 4'

            i += 1
            continue

        # Handle emphasis markers at start of line
        if line.strip().startswith('**') and not line.startswith('#'):
            # Bold emphasis line
            text = re.sub(r'\*\*', '', line).strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '黑体', 11, True, COLOR_DARK_GRAY)
            i += 1
            continue

        # Handle list items
        if re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line).strip()
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                set_chinese_font(run, '宋体', 10, False, COLOR_DARK_GRAY)
            i += 1
            continue

        # Handle numbered list
        if re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line).strip()
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                set_chinese_font(run, '宋体', 10, False, COLOR_DARK_GRAY)
            i += 1
            continue

        # Handle confidence/source notes inline
        if '[置信度:' in line or '[来源:' in line or '[来源：' in line:
            p = doc.add_paragraph()
            # Extract main text and note
            confidence_match = re.search(r'\[置信度:\s*([^\]]+)\]', line)
            source_match = re.search(r'\[来源:\s*([^\]]+)\]|\[来源：\s*([^\]]+)\]', line)

            main_text = line
            if confidence_match:
                main_text = main_text.replace(confidence_match.group(0), '').strip()
            if source_match:
                main_text = main_text.replace(source_match.group(0), '').strip()

            if main_text:
                run = p.add_run(main_text)
                set_chinese_font(run, '宋体', 10, False, COLOR_DARK_GRAY)

            if confidence_match:
                run = p.add_run(f' [置信度: {confidence_match.group(1).strip()}]')
                set_chinese_font(run, '宋体', 9, False, COLOR_ORANGE)

            p = doc.add_paragraph()
            if source_match:
                source_text = source_match.group(1) if source_match.group(1) else source_match.group(2)
                run = p.add_run(f'来源: {source_text.strip()}')
                set_chinese_font(run, '宋体', 8, False, RGBColor(100, 100, 100))

            i += 1
            continue

        # Handle blockquotes
        if line.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', line).strip()
            p = doc.add_paragraph(text, style='Quote')
            for run in p.runs:
                set_chinese_font(run, '宋体', 10, False, RGBColor(80, 80, 80))
            i += 1
            continue

        # Handle empty lines
        if not line.strip():
            i += 1
            continue

        # Handle regular paragraphs
        if line.strip():
            p = doc.add_paragraph(line.strip())
            for run in p.runs:
                set_chinese_font(run, '宋体', 11, False, COLOR_DARK_GRAY)
            i += 1
            continue

        i += 1

    # Save the document
    doc.save(output_path)
    print(f"✅ Word document generated: {output_path}")
    print(f"   Full report content converted successfully")


if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("Usage: python complete_markdown_to_word.py <input.md> [output.docx]")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
    else:
        output_path = input_path.with_suffix('.docx')

    convert_markdown_to_word(input_path, output_path)
